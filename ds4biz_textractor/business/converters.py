import email
import functools
import io
import logging
import os
import re
import time

import fitz
import pytesseract
import sanic
from PIL import Image
from PyPDF2 import PdfFileMerger
from bs4 import BeautifulSoup as bs
from docx import Document
from ds4biz_format_parsers.business.converters import Email2TextEmail

from ds4biz_textractor.business.ocr import TESSERACT
from ds4biz_textractor.config.app_config import PROCESS_WORKERS, DPI_DEFAULT
from ds4biz_textractor.dao.bb_dao import RDao
from ds4biz_textractor.utils.eml_utils import te2text
from ds4biz_textractor.utils.logger_utils import logger
from ds4biz_textractor.utils.pdf_utils import manipulate_pdf_page

os.environ['OMP_THREAD_LIMIT'] = '1'
from concurrent.futures import ProcessPoolExecutor
import asyncio

POOL = ProcessPoolExecutor(max_workers=PROCESS_WORKERS)


class PDF2text:
    async def __call__(self, file: sanic.request.File,
                       force_extraction: bool = False, **kwargs):
        print(file.name)
        logger.debug("force OCR extraction: %s" % str(force_extraction))
        configs = kwargs.get('configs')
        tesseract = TESSERACT(**configs)
        dpi = DPI_DEFAULT
        if "preprocessing_configs" in configs:
            dpi = configs.get("preprocessing_configs").get("dpi")
        loop = asyncio.get_event_loop()
        content = file.body
        doc = fitz.open("pdf", content)
        logger.debug(doc)
        for i, page in enumerate(doc):
            if force_extraction:
                text = ''
            else:
                text = page.get_text()
            logger.info(f"Pagina {i}")
            try:
                if len(re.sub(r'\s', '', str(text))) < 500:  # da passare a tesseract
                    text = text + '\n' if text else ''
                    img = manipulate_pdf_page(page, dpi)
                    text += await loop.run_in_executor(POOL, functools.partial(tesseract), img)
                    # text = tesseract(img)
                    yield dict(page=i, text=text, filename=file.name)
                else:
                    logger.debug(msg="machine readable file... getting text...")
                    yield dict(page=i, text=text, filename=file.name)
            except Exception as inst:
                logging.exception(inst)
                if text:
                    yield dict(page=i, text=text, filename=file.name)


class EML2text:

    def __init__(self):
        self.e2t = Email2TextEmail(None)

    async def __call__(self, file: sanic.request.File, **kwargs):
        try:
            content = file.body
            if hasattr(content, "read"):
                text = self.e2t(email.message_from_binary_file(content))
            else:
                te = self.e2t(email.message_from_bytes(content))
                text = te2text(te)
            yield dict(text=text, filename=file.name)
        except Exception as inst:
            logger.exception(inst)


class IMG2text:
    "Extract text from image formats: png jpg jpeg tif"

    async def __call__(self, file: sanic.request.File, **kwargs):
        configs = kwargs.get('configs')
        tesseract = TESSERACT(**configs)

        try:
            content = file.body
            if hasattr(content, "read"):

                text = tesseract(Image.open(content))
            else:
                text = tesseract(Image.open(io.BytesIO(content)))
            yield dict(text=text, filename=file.name)
        except Exception as inst:
            logger.exception(inst)


class TXT2text:

    async def __call__(self, file: sanic.request.File, **kwargs):
        try:
            # text = textract.process(fname)
            content = file.body
            if hasattr(content, "read"):
                text = content.read().decode('ISO-8859-1')
            else:
                text = content.decode('ISO-8859-1')
            yield dict(text=text, filename=file.name)
        except Exception as inst:
            logger.exception(inst)


class HTML2text:

    def __call__(self, content, **kwargs):
        raise Exception("Not Implemented")


class DOC2text:

    async def __call__(self, file: sanic.request.File, **kwargs):
        content = file.body
        if hasattr(content, "read"):
            text = self.convert(content.read())
        else:
            text = self.convert(content)
        yield dict(text=text, filename=file.name)

    def convert(self, content):
        soup = bs(content, features="lxml")
        [s.extract() for s in soup(['script'])]
        tmpText = soup.get_text()
        tmpText = re.split('\n', tmpText)
        text = []
        for row in tmpText:
            row = re.sub('(?<=[^ ]) (?=[^ ])', '', row)
            row = re.sub(' +', ' ', row)
            u = len(re.findall('(?i)[a-z0-9 ]', row))
            if len(row) > 3 and u / len(row) > .9:
                text.append(row.strip())
        return '\n'.join(text)


class DOCX2text:

    async def __call__(self, file: sanic.request.File, **kwargs):
        try:
            content = file.body
            # text = textract.process(fname)
            if hasattr(content, "read"):
                source_stream = io.StringIO(content.read())
                try:
                    doc = fitz.Document(source_stream)
                except Exception as insta:
                    logger.debug(insta.__dict__)
                p_texts = []
                for p in doc.paragraphs:
                    p_texts.append(p.text)
                text = '\n'.join(p_texts)
            else:
                source_stream = io.BytesIO(content)
                try:
                    doc = Document(source_stream)
                except Exception as insta:
                    logger.debug(insta.__dict__)
                p_texts = []
                for p in doc.paragraphs:
                    p_texts.append(p.text)
                text = '\n'.join(p_texts)
                # text = content.decode()
            yield dict(text=text, filename=file.name)
        except Exception as inst:
            logger.exception(inst)


class DOCDummyConverter:

    def extract(self, tmpText):
        tmpText = re.split('\n', tmpText)
        text = []
        for row in tmpText:
            ### lunghezza media dei token (se e' splittato per ogni lettera)
            tmean = [len(t) for t in row.split()]
            if tmean:
                tmean = sum(tmean) / len(tmean)
                if tmean < 3:
                    row = re.sub(r'(?<=[^ ]) (?=[^ ])', '', row)
                row = re.sub(' +', ' ', row)
                ### percentuale caratteri sensati
                u = len(re.findall(r'(?i)[a-z0-9\s.()/,;:\x00]', row))
                if len(row) > 3 and u / len(row) > .9:
                    text.append(row.strip())
        return '\n'.join(text)

    async def __call__(self, file: sanic.request.File, **kwargs):
        try:
            content = file.body
            # text = textract.process(fname)
            if hasattr(content, "read"):
                text = content.read().decode('latin')
            else:
                text = content.decode('latin')
            text = self.extract(text)
            yield dict(text=text, filename=file.name)
        except Exception as inst:
            logger.exception(inst)


#
# class HOCR:
#
#     def __init__(self, output: {"html", "pdf", "json"} = "json"):
#         self.output = output
#
#     def __call__(self, file, ct):
#         logger.debug("start HOCR process")
#         if "image/" in ct:
#             img = Image.open(io.BytesIO(file))
#             return self.img2hocr(img)
#         elif ct == "application/pdf":
#             return self.pdf2hocr(file)
#         else:
#             raise Exception("Not supported format")


class PDF2hocr:
    async def __call__(self, file: sanic.request.File,
                       output, **kwargs):
        logger.debug("pdf to HOCR")

        if output == "application/pdf":
            ret_pdf = PdfFileMerger()
        else:
            hocr_output = dict()

        configs = kwargs.get('configs')
        configs["hocr"] = True
        configs["output"] = output
        dpi = None
        if "preprocessing_configs" in configs:
            dpi = configs.get("preprocessing_configs").get("dpi")

        tesseract = TESSERACT(**configs)
        loop = asyncio.get_event_loop()
        content = file.body
        doc = fitz.open("pdf", content)
        logger.debug(doc)

        for i, page in enumerate(doc):
            logger.info(f"Pagina {i}")

            img = manipulate_pdf_page(page, dpi)
            res = await loop.run_in_executor(POOL, functools.partial(tesseract), img)
            if output == "application/pdf":
                ret_pdf.append(res)
            else:
                hocr_output[str(i)] = res
        logger.debug("pdf hocr done...")

        if output == "application/pdf":
            buffer = io.BytesIO()
            ret_pdf.write(buffer)
            buffer.seek(0)
            return buffer

        return hocr_output


class IMG2hocr:
    async def __call__(self, file: sanic.request.File,
                       output, **kwargs):

        configs = kwargs.get('configs')
        configs["hocr"] = True
        configs["output"] = output
        tesseract = TESSERACT(**configs)
        logger.debug("config hocr ----", configs)

        try:
            content = file.body
            if hasattr(content, "read"):

                res = tesseract(Image.open(content))
            else:
                res = tesseract(Image.open(io.BytesIO(content)))
            if output == "application/json":
                return res  # dict(content=res, filename=file.name)
            return res
        except Exception as inst:
            logger.exception(inst)


class ConverterFactory:

    def __init__(self, mapping=None):
        self.mapping = mapping or dict()

    def register(self, ext, value):
        self.mapping[ext] = value

    def get(self, ext):
        print("==========", ext)
        c = dict(jpg="img", jpeg="img", png="img", tif="img", tiff="img")
        if ext.startswith(hocr_f):
            temp_ext = ext.replace(hocr_f, "")
            if temp_ext in c:
                ext = hocr_f + c[temp_ext]
        if ext in c:
            ext = c[ext]
        if ext in self.mapping:
            return self.mapping[ext]
        print(ext)
        raise Exception('Extension not handled: {ext}'.format(ext=ext))
    #
    # try:
    #     temp_file = save2temp(file)
    #     ext = temp_file.name.split('.')[-1]
    #     if ext in HANDLER_MAPPING:
    #         for el in file_handler(temp_file):
    #             output = f2t(el['file'].name)
    #             ti.text += '\n'+output
    #         return ti
    #     else:
    #         output = f2t(temp_file.name)
    #         ti.text = output
    #         return ti


CONV_FACTORY = ConverterFactory()

CONV_FACTORY.register("pdf", PDF2text)
CONV_FACTORY.register("p7m", PDF2text)
CONV_FACTORY.register("eml", EML2text)
CONV_FACTORY.register("img", IMG2text)
CONV_FACTORY.register("docx", DOCX2text)
CONV_FACTORY.register('doc', DOCDummyConverter)
CONV_FACTORY.register('txt', TXT2text)

hocr_f = "hocr_"
CONV_FACTORY.register(hocr_f + "pdf", PDF2hocr)
# CONV_FACTORY.register(hocr_f + "p7m", PDF2hocr)
# CONV_FACTORY.register(hocr_f + "eml", EML2text)
CONV_FACTORY.register(hocr_f + "img", IMG2hocr)
# CONV_FACTORY.register(hocr_f + "docx", DOCX2text)
# CONV_FACTORY.register(hocr_f + 'doc', DOCDummyConverter)
# CONV_FACTORY.register(hocr_f + 'txt', TXT2text)
